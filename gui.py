#!/usr/bin/env python3
"""
Receipt Generator — GUI front-end
Launch with:  python gui.py
"""

import os
import re
import subprocess
import sys
import threading
from datetime import datetime

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox

# ── Import shared config + helpers from main.py ──────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)

from main import (
    TEMPLATE_PATH, OUTPUT_FOLDER, FOXIT_PATH, EXCEL_PATH,
    RECEIPT_NUM_LABEL, DATE_LABEL, MONTH_LABEL, SUM_LABEL,
    HEBREW_MONTHS, SUM_LABEL_ALTERNATIVES,
    _word_replace_all, _find_label_sep_value, _update_field,
    _get_amount_from_excel,
)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _read_next_receipt_num() -> int:
    """Read the current receipt number from the Word template and return N+1."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed.  Run:  pip install python-docx")

    doc = Document(TEMPLATE_PATH)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    full = '\n'.join(parts)
    _, num_str = _find_label_sep_value(full, RECEIPT_NUM_LABEL, r'\d+')
    return int(num_str) + 1 if num_str else 1


def _produce_receipt(month_he: str, amount: str, new_num: int, log_fn):
    """
    Full receipt flow: update Word → export PDF → run place_sig_once.py.
    log_fn(msg) is called with progress messages (safe to call from any thread).
    """
    import win32com.client

    month_en = HEBREW_MONTHS[month_he]
    today    = datetime.now()
    date_str = today.strftime('%d/%m/%Y')
    year     = today.year

    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    abs_template = os.path.abspath(TEMPLATE_PATH)

    try:
        doc      = word.Documents.Open(abs_template)
        doc_text = doc.Content.Text
        content  = doc.Content

        log_fn('Updating document fields...')

        # Receipt number — replace old value with the one the user confirmed
        sep_r, old_num_str = _find_label_sep_value(doc_text, RECEIPT_NUM_LABEL, r'\d+')
        if not old_num_str:
            raise RuntimeError(f"'{RECEIPT_NUM_LABEL}' not found in template.")
        _word_replace_all(
            content,
            RECEIPT_NUM_LABEL + sep_r + old_num_str,
            RECEIPT_NUM_LABEL + sep_r + str(new_num),
        )
        log_fn(f'  ✓  Receipt #: {old_num_str} → {new_num}')

        _update_field(content, doc_text, DATE_LABEL, date_str,
                      r'\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}')
        _update_field(content, doc_text, MONTH_LABEL, month_he,
                      '|'.join(re.escape(m) for m in HEBREW_MONTHS))
        _update_field(content, doc_text, SUM_LABEL, amount,
                      r'[\d,\.]+',
                      label_alternatives=SUM_LABEL_ALTERNATIVES[1:])

        doc.Save()
        log_fn('✓  Document saved')

        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        pdf_filename = f'recipt_{new_num}_{month_en}_{year}.pdf'
        pdf_path     = os.path.abspath(os.path.join(OUTPUT_FOLDER, pdf_filename))
        doc.ExportAsFixedFormat(pdf_path, 17)
        log_fn(f'✓  PDF saved: {pdf_filename}')
        doc.Close(False)

    finally:
        try:
            word.Quit()
        except Exception:
            pass

    if not os.path.exists(FOXIT_PATH):
        raise RuntimeError(f'Foxit not found:\n  {FOXIT_PATH}')

    sig_script = os.path.join(_HERE, 'place_sig_once.py')
    log_fn('Applying signature in Foxit...')
    result = subprocess.run([sys.executable, sig_script, pdf_path])
    if result.returncode != 0:
        raise RuntimeError('Signature script exited with an error.')
    log_fn('✓  All done!')


# ── GUI ───────────────────────────────────────────────────────────────────────

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('Receipt Generator')
        self.resizable(False, False)
        self._build_ui()
        # Pre-fill receipt number (month not chosen yet, amount will fill on selection)
        self._refresh_receipt_num()

    def _build_ui(self):
        pad = {'padx': 14, 'pady': 7}

        # ── Month ──────────────────────────────────────────────────────────
        tk.Label(self, text='Month:', anchor='w', width=14).grid(
            row=0, column=0, sticky='w', **pad)
        self.month_var = tk.StringVar()
        months = list(HEBREW_MONTHS.keys())
        self.month_cb = ttk.Combobox(
            self, textvariable=self.month_var,
            values=months, state='readonly', width=22)
        self.month_cb.grid(row=0, column=1, sticky='ew', **pad)
        self.month_cb.bind('<<ComboboxSelected>>', lambda e: self._on_month_change())

        # ── Receipt # ──────────────────────────────────────────────────────
        tk.Label(self, text='Receipt #:', anchor='w', width=14).grid(
            row=1, column=0, sticky='w', **pad)
        self.num_var = tk.StringVar()
        tk.Entry(self, textvariable=self.num_var, width=24).grid(
            row=1, column=1, sticky='ew', **pad)

        # ── Amount ─────────────────────────────────────────────────────────
        tk.Label(self, text='Amount:', anchor='w', width=14).grid(
            row=2, column=0, sticky='w', **pad)
        self.amount_var = tk.StringVar()
        tk.Entry(self, textvariable=self.amount_var, width=24).grid(
            row=2, column=1, sticky='ew', **pad)

        # ── Button ─────────────────────────────────────────────────────────
        self.btn = tk.Button(
            self, text='  Produce Receipt  ',
            command=self._run,
            bg='#2e7d32', fg='white',
            font=('Segoe UI', 11, 'bold'),
            relief='flat', padx=8, pady=8,
            cursor='hand2',
        )
        self.btn.grid(row=3, column=0, columnspan=2, pady=12)

        # ── Log ────────────────────────────────────────────────────────────
        self.log = scrolledtext.ScrolledText(
            self, height=13, width=58,
            state='disabled', font=('Consolas', 9),
            bg='#1e1e1e', fg='#d4d4d4',
            insertbackground='white',
        )
        self.log.grid(row=4, column=0, columnspan=2, padx=14, pady=(0, 14))

    # ── Event handlers ────────────────────────────────────────────────────────

    def _refresh_receipt_num(self):
        try:
            n = _read_next_receipt_num()
            self.num_var.set(str(n))
        except Exception as e:
            self._log(f'(Could not read receipt #: {e})')

    def _on_month_change(self):
        month_he = self.month_var.get()
        if not month_he:
            return
        month_en = HEBREW_MONTHS[month_he]
        try:
            amt = _get_amount_from_excel(month_en)
            self.amount_var.set(amt)
        except Exception as e:
            self.amount_var.set('')
            self._log(f'(Could not read amount from Excel: {e})')

    def _log(self, msg: str):
        """Thread-safe log append."""
        def _append():
            self.log.config(state='normal')
            self.log.insert('end', msg + '\n')
            self.log.see('end')
            self.log.config(state='disabled')
        self.after(0, _append)

    def _run(self):
        month_he = self.month_var.get()
        if not month_he:
            messagebox.showwarning('Missing input', 'Please select a month.')
            return
        amount = self.amount_var.get().strip()
        if not amount:
            messagebox.showwarning('Missing input', 'Please enter an amount.')
            return
        try:
            new_num = int(self.num_var.get().strip())
        except ValueError:
            messagebox.showwarning('Invalid input', 'Receipt # must be a whole number.')
            return

        self.btn.config(state='disabled')
        self._log(f'\n── Receipt #{new_num} | {month_he} | ₪{amount} ──')

        def job():
            try:
                _produce_receipt(month_he, amount, new_num, self._log)
            except Exception as exc:
                self._log(f'ERROR: {exc}')
            finally:
                self.after(0, lambda: self.btn.config(state='normal'))
                # Refresh receipt number for the next run
                self.after(500, self._refresh_receipt_num)

        threading.Thread(target=job, daemon=True).start()


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    App().mainloop()

